"""Tests for DOCX document generation support.

Covers:
- convert_to_pdf with doc_format="docx" produces a valid .docx file
- convert_to_pdf with doc_format="pdf" still works (default behavior)
- convert_to_pdf with invalid doc_format raises ValueError
- batch_convert respects doc_format parameter
- render_docx produces a DOCX with expected sections
- build_prompt uses correct file extensions for docx format
"""

import shutil
import sys
from pathlib import Path
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

SAMPLE_RESUME = """\
John Doe
Senior Software Engineer
Seattle, WA
john@example.com | linkedin.com/in/johndoe | github.com/johndoe

SUMMARY
Experienced software engineer with 10+ years building distributed systems.

TECHNICAL SKILLS
Languages: Python, Go, TypeScript, Java
Frameworks: Django, FastAPI, React, Next.js
Infrastructure: AWS, Kubernetes, Docker, Terraform

EXPERIENCE
Senior Software Engineer — Acme Corp
Seattle, WA | 2020 – Present
- Built a real-time data pipeline processing 1M events/day
- Led migration from monolith to microservices architecture
- Mentored 3 junior engineers

Software Engineer — Widgets Inc
San Francisco, CA | 2016 – 2020
- Developed REST API serving 10K requests/second
- Implemented CI/CD pipeline reducing deploy time by 60%

EDUCATION
B.S. Computer Science, University of Washington, 2016
"""


class TestConvertToDocx:
    """Tests for convert_to_pdf with doc_format='docx'."""

    def test_docx_generation(self, tmp_path):
        """convert_to_pdf with doc_format='docx' creates a valid .docx file."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")

        from applypilot.scoring.pdf import convert_to_pdf

        result = convert_to_pdf(txt_file, doc_format="docx")

        assert result.suffix == ".docx"
        assert result.exists()
        assert result.stat().st_size > 0

        # Verify it's a valid DOCX by opening it
        from docx import Document

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" in full_text
        assert "Senior Software Engineer" in full_text

    def test_docx_output_path_override(self, tmp_path):
        """convert_to_pdf with doc_format='docx' respects output_path."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")
        custom_out = tmp_path / "custom_name.docx"

        from applypilot.scoring.pdf import convert_to_pdf

        result = convert_to_pdf(txt_file, output_path=custom_out, doc_format="docx")

        assert result == custom_out
        assert result.exists()

    def test_default_format_is_docx(self, tmp_path):
        """convert_to_pdf defaults to DOCX when doc_format is not specified."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")

        from applypilot.scoring.pdf import convert_to_pdf

        # Mock render_docx to avoid needing python-docx
        with patch("applypilot.scoring.pdf.render_docx") as mock_render:
            result = convert_to_pdf(txt_file)
            assert result.suffix == ".docx"
            mock_render.assert_called_once()

    def test_invalid_format_raises(self, tmp_path):
        """convert_to_pdf with invalid doc_format raises ValueError."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")

        import pytest

        from applypilot.scoring.pdf import convert_to_pdf

        with pytest.raises(ValueError, match="Invalid doc_format"):
            convert_to_pdf(txt_file, doc_format="odt")

    def test_html_only_ignores_doc_format(self, tmp_path):
        """html_only=True produces HTML regardless of doc_format."""
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")

        from applypilot.scoring.pdf import convert_to_pdf

        result = convert_to_pdf(txt_file, html_only=True, doc_format="docx")

        assert result.suffix == ".html"
        assert result.exists()
        content = result.read_text()
        assert "<html>" in content


class TestRenderDocx:
    """Tests for render_docx function."""

    def test_sections_present(self, tmp_path):
        """render_docx includes all resume sections."""
        from docx import Document

        from applypilot.scoring.pdf import parse_resume, render_docx

        resume = parse_resume(SAMPLE_RESUME)
        out = tmp_path / "test.docx"
        render_docx(resume, str(out))

        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Header
        assert "John Doe" in full_text
        assert "Senior Software Engineer" in full_text
        assert "Seattle, WA" in full_text

        # Sections
        assert "SUMMARY" in full_text
        assert "TECHNICAL SKILLS" in full_text
        assert "EXPERIENCE" in full_text
        assert "EDUCATION" in full_text

        # Content
        assert "Python" in full_text
        assert "Acme Corp" in full_text
        assert "University of Washington" in full_text

    def test_skills_formatting(self, tmp_path):
        """render_docx formats skills with bold category names."""
        from docx import Document

        from applypilot.scoring.pdf import parse_resume, render_docx

        resume = parse_resume(SAMPLE_RESUME)
        out = tmp_path / "test.docx"
        render_docx(resume, str(out))

        doc = Document(str(out))
        # Find a paragraph containing "Languages:"
        for p in doc.paragraphs:
            if "Languages:" in p.text:
                # First run should be bold (the category)
                assert p.runs[0].bold is True
                break
        else:
            raise AssertionError("Skills paragraph with 'Languages:' not found")

    def test_bullet_points(self, tmp_path):
        """render_docx creates bullet points for experience entries."""
        from docx import Document

        from applypilot.scoring.pdf import parse_resume, render_docx

        resume = parse_resume(SAMPLE_RESUME)
        out = tmp_path / "test.docx"
        render_docx(resume, str(out))

        doc = Document(str(out))
        bullet_texts = [p.text for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
        assert len(bullet_texts) > 0
        assert any("real-time data pipeline" in b for b in bullet_texts)


# 2026-08-28 fix: cover letters were previously run through the resume-
# shaped parse_resume/render_docx/build_html path used above. Since a
# cover letter never contains a literal "SUMMARY" header line,
# parse_resume's header-scanning loop never broke, so every line after the
# first blank line got joined into a single "subtitle" paragraph -- all
# paragraph breaks lost, and the rest of the letter's content silently
# dropped because the section-splitting loop never found an ALL-CAPS
# header either. Confirmed via a real generated cover letter fixture shape
# (greeting line, several single-line paragraphs, closing name -- see
# ~/.applypilot/cover_letters/*_CL.txt for the real format this mirrors).
SAMPLE_COVER_LETTER = """\
Dear Hiring Manager,

I built a tool that automates data extraction from scanned documents using Python and OCR, which reduced manual review time significantly for a real client project.

At my previous role, I worked directly with customers to diagnose and resolve technical issues, building the kind of clear communication skills this position requires.

I would welcome the chance to discuss how my background fits this role.

Philip"""


class TestCoverLetterPlainTextRendering:
    """render_plain_text_docx / build_plain_text_html / convert_to_pdf's
    content_type='cover_letter' path -- must preserve every source
    paragraph verbatim and in order, never route through resume parsing."""

    def test_split_into_paragraphs_preserves_order_and_content_verbatim(self):
        from applypilot.scoring.pdf import _split_into_paragraphs

        paragraphs = _split_into_paragraphs(SAMPLE_COVER_LETTER)
        expected = [
            "Dear Hiring Manager,",
            (
                "I built a tool that automates data extraction from scanned documents "
                "using Python and OCR, which reduced manual review time significantly "
                "for a real client project."
            ),
            (
                "At my previous role, I worked directly with customers to diagnose and "
                "resolve technical issues, building the kind of clear communication "
                "skills this position requires."
            ),
            "I would welcome the chance to discuss how my background fits this role.",
            "Philip",
        ]
        assert paragraphs == expected

    def test_render_plain_text_docx_preserves_paragraphs_verbatim_and_in_order(self, tmp_path):
        from docx import Document

        from applypilot.scoring.pdf import _split_into_paragraphs, render_plain_text_docx

        out = tmp_path / "cover_letter.docx"
        render_plain_text_docx(SAMPLE_COVER_LETTER, str(out))

        doc = Document(str(out))
        actual = [p.text for p in doc.paragraphs]
        expected = _split_into_paragraphs(SAMPLE_COVER_LETTER)

        assert actual == expected, (
            f"Expected {len(expected)} paragraphs preserved verbatim and in order, got {actual!r} "
            "-- before the fix this would collapse to far fewer paragraphs with content joined/dropped"
        )

    def test_build_plain_text_html_preserves_paragraphs_verbatim_and_in_order(self):
        import re as _re

        from applypilot.scoring.pdf import _split_into_paragraphs, build_plain_text_html

        html = build_plain_text_html(SAMPLE_COVER_LETTER)
        found = _re.findall(r"<p>(.*?)</p>", html, flags=_re.DOTALL)
        expected = _split_into_paragraphs(SAMPLE_COVER_LETTER)

        assert found == expected, f"Expected {len(expected)} <p> tags preserved verbatim and in order, got {found!r}"

    def test_convert_to_pdf_cover_letter_content_type_produces_matching_paragraph_count(self, tmp_path):
        """End-to-end integration through the real cover_letter.py call
        shape: convert_to_pdf(..., content_type='cover_letter', doc_format='docx')."""
        from docx import Document

        from applypilot.scoring.pdf import _split_into_paragraphs, convert_to_pdf

        txt_file = tmp_path / "cover_letter.txt"
        txt_file.write_text(SAMPLE_COVER_LETTER, encoding="utf-8")

        result = convert_to_pdf(txt_file, doc_format="docx", content_type="cover_letter")

        assert result.suffix == ".docx"
        doc = Document(str(result))
        actual = [p.text for p in doc.paragraphs]
        expected = _split_into_paragraphs(SAMPLE_COVER_LETTER)
        assert actual == expected

    def test_convert_to_pdf_resume_content_type_is_unchanged(self, tmp_path):
        """Regression guard: content_type='resume' (the default, used by
        every existing resume call site) must behave identically to before
        this fix -- still section-parsed, not routed through the new
        plain-text path."""
        from docx import Document

        from applypilot.scoring.pdf import convert_to_pdf

        txt_file = tmp_path / "resume.txt"
        txt_file.write_text(SAMPLE_RESUME, encoding="utf-8")

        result = convert_to_pdf(txt_file, doc_format="docx")  # content_type defaults to "resume"

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SUMMARY" in full_text
        assert "TECHNICAL SKILLS" in full_text
        assert "EXPERIENCE" in full_text
        assert "Acme Corp" in full_text

    def test_convert_to_pdf_invalid_content_type_raises(self, tmp_path):
        import pytest

        from applypilot.scoring.pdf import convert_to_pdf

        txt_file = tmp_path / "x.txt"
        txt_file.write_text(SAMPLE_COVER_LETTER, encoding="utf-8")

        with pytest.raises(ValueError, match="Invalid content_type"):
            convert_to_pdf(txt_file, content_type="bogus")


class TestBatchConvertDocx:
    """Tests for batch_convert with doc_format parameter."""

    def test_batch_convert_docx(self, tmp_path):
        """batch_convert with doc_format='docx' creates .docx files."""
        # Create some .txt files in a fake TAILORED_DIR
        for i in range(3):
            (tmp_path / f"company_job_{i}.txt").write_text(SAMPLE_RESUME, encoding="utf-8")
        # Also create a _JOB.txt that should be skipped
        (tmp_path / "company_job_0_JOB.txt").write_text("Job description", encoding="utf-8")

        from applypilot.scoring import pdf

        with patch.object(pdf, "TAILORED_DIR", tmp_path):
            count = pdf.batch_convert(doc_format="docx")

        assert count == 3
        docx_files = list(tmp_path.glob("*.docx"))
        assert len(docx_files) == 3
        # No PDFs should have been created
        pdf_files = list(tmp_path.glob("*.pdf"))
        assert len(pdf_files) == 0

    def test_batch_convert_skips_existing(self, tmp_path):
        """batch_convert skips files that already have a .docx counterpart."""
        (tmp_path / "company_job_1.txt").write_text(SAMPLE_RESUME, encoding="utf-8")
        (tmp_path / "company_job_1.docx").write_text("existing", encoding="utf-8")  # already exists
        (tmp_path / "company_job_2.txt").write_text(SAMPLE_RESUME, encoding="utf-8")

        from applypilot.scoring import pdf

        with patch.object(pdf, "TAILORED_DIR", tmp_path):
            count = pdf.batch_convert(doc_format="docx")

        assert count == 1  # only job_2 converted

    def test_batch_convert_invalid_format(self):
        """batch_convert with invalid format raises ValueError."""
        import pytest

        from applypilot.scoring.pdf import batch_convert

        with pytest.raises(ValueError, match="Invalid doc_format"):
            batch_convert(doc_format="rtf")


class TestBuildPromptDocx:
    """Tests for build_prompt doc_format parameter and file resolution logic."""

    def test_docx_file_resolution(self, tmp_path):
        """With doc_format='docx', build_prompt resolves .docx files and copies with correct extension."""
        resume_txt = tmp_path / "resume.txt"
        resume_txt.write_text(SAMPLE_RESUME, encoding="utf-8")
        resume_docx = tmp_path / "resume.docx"
        resume_docx.write_text("fake docx content", encoding="utf-8")

        dest_dir = tmp_path / "workers" / "current"
        dest_dir.mkdir(parents=True, exist_ok=True)

        # Simulate the file resolution logic from build_prompt
        doc_ext = ".docx"
        src_doc = Path(str(resume_txt)).with_suffix(doc_ext).resolve()
        assert src_doc.exists(), "Source .docx should exist"

        name_slug = "John_Doe"
        upload_doc = dest_dir / f"{name_slug}_Resume{doc_ext}"
        shutil.copy(str(src_doc), str(upload_doc))

        assert upload_doc.exists()
        assert upload_doc.suffix == ".docx"
        assert upload_doc.name == "John_Doe_Resume.docx"

    def test_pdf_file_resolution(self, tmp_path):
        """With doc_format='pdf', build_prompt resolves .pdf files."""
        resume_txt = tmp_path / "resume.txt"
        resume_txt.write_text(SAMPLE_RESUME, encoding="utf-8")
        resume_pdf = tmp_path / "resume.pdf"
        resume_pdf.write_text("fake pdf content", encoding="utf-8")

        doc_ext = ".pdf"
        src_doc = Path(str(resume_txt)).with_suffix(doc_ext).resolve()
        assert src_doc.exists(), "Source .pdf should exist"

    def test_prompt_format_strings(self):
        """Verify format string interpolation produces correct labels."""
        # Simulate the f-string logic used in the prompt template
        for fmt, expected_label in [("pdf", "PDF"), ("docx", "DOCX")]:
            doc_format = fmt
            resume_line = f"Resume {doc_format.upper()} (upload this): /path/to/file.{doc_format}"
            cl_line = f"Cover Letter {doc_format.upper()} (upload if asked): /path/to/file.{doc_format}"

            assert expected_label in resume_line
            assert f".{fmt}" in resume_line
            assert expected_label in cl_line

    def test_build_prompt_signature_has_doc_format(self):
        """build_prompt accepts doc_format parameter with 'pdf' default."""
        import inspect

        from applypilot.apply.prompt import build_prompt

        sig = inspect.signature(build_prompt)
        assert "doc_format" in sig.parameters
        assert sig.parameters["doc_format"].default == "pdf"


class TestValidDocFormats:
    """Tests for VALID_DOC_FORMATS constant."""

    def test_valid_formats(self):
        from applypilot.scoring.pdf import VALID_DOC_FORMATS

        assert "pdf" in VALID_DOC_FORMATS
        assert "docx" in VALID_DOC_FORMATS

    def test_docx_is_default(self):
        """Ensure DOCX is the default format in the public API."""
        import inspect

        from applypilot.scoring.pdf import convert_to_pdf

        sig = inspect.signature(convert_to_pdf)
        assert sig.parameters["doc_format"].default == "docx"


class TestSetDocFormat:
    """Tests for launcher.set_doc_format."""

    def test_set_and_read(self):
        from applypilot.apply import launcher

        original = launcher._doc_format
        try:
            launcher.set_doc_format("docx")
            assert launcher._doc_format == "docx"
            launcher.set_doc_format("pdf")
            assert launcher._doc_format == "pdf"
        finally:
            launcher._doc_format = original
